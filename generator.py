import copy
import io
import os
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

TEMPLATE_PATH = os.path.join(os.path.dirname(__file__), "근로계약서_양식.docx")

CHECK = "■"
UNCHECK = "□"


def _get_para_text(para):
    return "".join(r.text for r in para.runs)


def _set_cell_text(cell, text):
    """셀의 텍스트를 교체하되 첫 번째 run의 서식을 유지한다."""
    para = cell.paragraphs[0]
    rpr_xml = None
    if para.runs:
        rpr_xml = para.runs[0]._r.find(qn("w:rPr"))

    for child in list(para._p):
        if child.tag in (qn("w:r"), qn("w:hyperlink")):
            para._p.remove(child)

    if not text:
        return

    r = OxmlElement("w:r")
    if rpr_xml is not None:
        r.append(copy.deepcopy(rpr_xml))
    t = OxmlElement("w:t")
    t.text = text
    t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    r.append(t)
    para._p.append(r)


def _rewrite_para(para, new_text):
    """단락의 모든 run을 지우고 첫 번째 run 서식으로 새 텍스트를 쓴다."""
    rpr_xml = None
    if para.runs:
        rpr_xml = para.runs[0]._r.find(qn("w:rPr"))

    for child in list(para._p):
        if child.tag == qn("w:r"):
            para._p.remove(child)

    r = OxmlElement("w:r")
    if rpr_xml is not None:
        r.append(copy.deepcopy(rpr_xml))
    t = OxmlElement("w:t")
    t.text = new_text
    t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    r.append(t)
    para._p.append(r)


def generate_contract(data: dict) -> bytes:
    doc = Document(TEMPLATE_PATH)
    tables = doc.tables

    # ── Table 0: 계약 당사자 ──────────────────────────────────────
    # Row 0: 사업주 성명(0) | 값(1) | 사업자등록번호(2) | 값(3)
    # Row 3: 근로자 성명(0) | 값(1) | 주민등록번호(2) | 값(3)
    # Row 4: 근로자 주소(0) | 값(1 merged)
    # Row 5: 연락처(0) | 값(1) | 관계(2) | 값(3)
    t0 = tables[0]
    _set_cell_text(t0.rows[0].cells[1], data["owner_name"])
    _set_cell_text(t0.rows[0].cells[3], data["business_reg_no"])
    _set_cell_text(t0.rows[3].cells[1], data["worker_name"])
    _set_cell_text(t0.rows[3].cells[3], data["worker_id_no"])
    _set_cell_text(t0.rows[4].cells[1], data["worker_address"])
    _set_cell_text(t0.rows[5].cells[1], data["worker_phone"])
    _set_cell_text(t0.rows[5].cells[3], data["worker_relation"])

    # ── Table 1: 제1조 근무장소/담당업무 ─────────────────────────
    # Row 0: 근무장소(0) | 값(1)   Row 1: 담당업무(0) | 값(1)
    t1 = tables[1]
    _set_cell_text(t1.rows[1].cells[1], data["job_duties"])

    # ── Table 2: 제2조 계약기간/수습기간 ─────────────────────────
    t2 = tables[2]
    # 계약기간
    if data["no_fixed_term"]:
        period_text = (
            f"{data['contract_start']} 부터  {data['contract_end']} 까지"
            f"  {CHECK} 기간의 정함이 없음"
        )
    else:
        period_text = (
            f"{data['contract_start']} 부터  {data['contract_end']} 까지"
            f"  {UNCHECK} 기간의 정함이 없음"
        )
    _set_cell_text(t2.rows[0].cells[1], period_text)

    # 수습기간
    if data["probation_none"]:
        prob_text = f"{CHECK} 없음   {UNCHECK} 있음"
    else:
        prob_text = (
            f"{UNCHECK} 없음   {CHECK} 있음"
            f" ({data['probation_months']}개월, 수습 중 임금: {data['probation_wage']}원)"
        )
    _set_cell_text(t2.rows[1].cells[1], prob_text)

    # ── Table 3: 제3조 근로시간 ──────────────────────────────────
    t3 = tables[3]
    work_text = (
        f"{data['work_start']} ~ {data['work_end']}"
        f"  (1일 {data['daily_hours']}시간,  1주 {data['weekly_hours']}시간)"
    )
    _set_cell_text(t3.rows[0].cells[1], work_text)
    _set_cell_text(t3.rows[1].cells[1], f"{data['break_start']} ~ {data['break_end']}")
    _set_cell_text(
        t3.rows[2].cells[1],
        f"매주 {data['work_days_per_week']}일 근무 / 주휴일: 매주 {data['weekly_holiday']}요일",
    )

    # ── Table 4: 제4조 임금 ──────────────────────────────────────
    t4 = tables[4]
    _set_cell_text(
        t4.rows[0].cells[1],
        f"월  금 {data['basic_wage']}원  (시급: {data['hourly_wage']}원)",
    )
    _set_cell_text(
        t4.rows[1].cells[1],
        f"식대: {data['meal_allowance']}원 /  기타: {data['other_allowance']}원",
    )
    _set_cell_text(t4.rows[2].cells[1], f"월  금 {data['total_wage']}원")
    _set_cell_text(
        t4.rows[3].cells[1],
        f"매월 {data['pay_day']}일  (지급일이 휴일인 경우 전일 지급)",
    )
    d_mark = CHECK if data["pay_direct"] else UNCHECK
    b_mark = CHECK if data["pay_bank"] else UNCHECK
    pay_text = (
        f"{d_mark} 직접지급   {b_mark} 통장입금"
        f"  (은행명: {data['bank_name']}, 계좌번호: {data['account_no']})"
    )
    _set_cell_text(t4.rows[4].cells[1], pay_text)

    # ── 제6조 사회보험 (단락) ────────────────────────────────────
    ins_map = {
        "고용보험": data["ins_employment"],
        "산재보험": data["ins_industrial"],
        "국민연금": data["ins_pension"],
        "건강보험": data["ins_health"],
    }
    for para in doc.paragraphs:
        full = _get_para_text(para)
        if "고용보험" in full and "산재보험" in full:
            new = full
            for name, checked in ins_map.items():
                new = new.replace(f"{CHECK} {name}", f"__CHK_{name}__")
                new = new.replace(f"{UNCHECK} {name}", f"__CHK_{name}__")
                new = new.replace(f"__CHK_{name}__", f"{CHECK if checked else UNCHECK} {name}")
            _rewrite_para(para, new)
            break

    # ── Table 5: 제7조 퇴직금/특약 ──────────────────────────────
    t5 = tables[5]
    y = CHECK if data["severance_yes"] else UNCHECK
    n = CHECK if data["severance_no"] else UNCHECK
    _set_cell_text(
        t5.rows[0].cells[1],
        f"{y} 해당  {n} 미해당  (1주 15시간 이상, 1년 이상 근무 시 해당)",
    )
    _set_cell_text(t5.rows[1].cells[1], data["special_terms"])

    # ── 계약체결일 (단락) ────────────────────────────────────────
    for para in doc.paragraphs:
        if "계약체결일" in _get_para_text(para):
            _rewrite_para(
                para,
                f"계약체결일:  {data['contract_year']}년  {data['contract_month']}월  {data['contract_day']}일",
            )
            break

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
